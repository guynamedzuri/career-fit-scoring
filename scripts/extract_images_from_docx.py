#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
DOCX 파일에서 이미지를 추출하여 저장하는 스크립트

사용법:
    python3 scripts/extract_images_from_docx.py <docx_file> <output_dir>
"""

import sys
import os
import json
from pathlib import Path
from zipfile import ZipFile

# Windows에서 한글 경로 처리
if sys.platform == 'win32':
    import locale
    if sys.stdout.encoding != 'utf-8':
        try:
            sys.stdout.reconfigure(encoding='utf-8')
            sys.stderr.reconfigure(encoding='utf-8')
        except:
            pass

try:
    from docx import Document
except ImportError:
    error_msg = json.dumps({"error": "python-docx 라이브러리가 필요합니다.\n설치 방법: pip3 install python-docx"})
    print(error_msg, file=sys.stderr)
    sys.exit(1)


def extract_images_from_docx(docx_path: str, output_dir: str) -> dict:
    """
    DOCX 파일에서 이미지를 추출하여 저장
    
    Args:
        docx_path: DOCX 파일 경로
        output_dir: 이미지를 저장할 디렉토리 경로
    
    Returns:
        dict: 추출된 이미지 정보
    """
    docx_path = os.path.normpath(docx_path)
    output_dir = os.path.normpath(output_dir)
    
    if not os.path.exists(docx_path):
        return {
            "error": f"File not found: '{docx_path}'"
        }
    
    # 출력 디렉토리 생성
    os.makedirs(output_dir, exist_ok=True)
    
    try:
        # DOCX 파일을 ZIP으로 열어서 이미지 추출
        extracted_images = []
        
        with ZipFile(docx_path, 'r') as zip_ref:
            # word/media/ 폴더에 있는 모든 이미지 파일 찾기
            image_files = [f for f in zip_ref.namelist() if f.startswith('word/media/')]
            
            for img_path in image_files:
                # 파일명 추출
                filename = os.path.basename(img_path)
                # 확장자 확인
                ext = os.path.splitext(filename)[1].lower()
                
                # 이미지 파일만 처리
                if ext in ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff', '.tif']:
                    # 파일 읽기
                    img_data = zip_ref.read(img_path)
                    
                    # 출력 파일 경로
                    output_path = os.path.join(output_dir, filename)
                    
                    # 파일 저장
                    with open(output_path, 'wb') as f:
                        f.write(img_data)
                    
                    extracted_images.append({
                        "original_path": img_path,
                        "filename": filename,
                        "output_path": output_path,
                        "size": len(img_data)
                    })
        
        # python-docx로 문서 열어서 이미지 위치 정보도 추출
        doc = Document(docx_path)
        
        # 각 테이블의 각 셀에서 이미지 찾기
        image_positions = []
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    # 셀 내 이미지 확인
                    for para in cell.paragraphs:
                        for run in para.runs:
                            if run.element.findall('.//a:blip', namespaces={'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'}):
                                # 이미지가 있는 셀
                                image_positions.append({
                                    "table_index": table_idx,
                                    "row_index": row_idx,
                                    "cell_index": cell_idx
                                })
                                break
        
        return {
            "success": True,
            "docx_path": docx_path,
            "output_dir": output_dir,
            "extracted_count": len(extracted_images),
            "images": extracted_images,
            "image_positions": image_positions  # 이미지가 있는 셀 위치
        }
    
    except Exception as e:
        return {
            "error": f"Failed to extract images: {str(e)}"
        }


def main():
    if len(sys.argv) < 3:
        error_msg = json.dumps({"error": "Usage: python3 scripts/extract_images_from_docx.py <docx_file> <output_dir>"})
        print(error_msg, file=sys.stderr)
        sys.exit(1)
    
    docx_path = sys.argv[1]
    output_dir = sys.argv[2]
    
    result = extract_images_from_docx(docx_path, output_dir)
    
    # JSON 출력
    print(json.dumps(result, ensure_ascii=False, indent=2))
    
    if "error" in result:
        sys.exit(1)


if __name__ == "__main__":
    main()
