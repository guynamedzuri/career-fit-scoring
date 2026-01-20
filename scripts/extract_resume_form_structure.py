#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
이력서 양식(DOCX)에서 모든 테이블과 셀을 추출하여 위치 정보를 출력하는 스크립트

사용법:
    python3 scripts/extract_resume_form_structure.py resume_form.docx
"""

import sys
import json
import os
from pathlib import Path
from io import BytesIO

# Windows에서 한글 경로 처리
if sys.platform == 'win32':
    import locale
    # 콘솔 인코딩 설정
    if sys.stdout.encoding != 'utf-8':
        try:
            sys.stdout.reconfigure(encoding='utf-8')
            sys.stderr.reconfigure(encoding='utf-8')
        except:
            pass

try:
    from docx import Document
    from docx.table import Table
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.oxml.ns import qn
except ImportError:
    print("ERROR: python-docx 라이브러리가 필요합니다.")
    print("설치 방법: pip3 install python-docx")
    sys.exit(1)


def extract_table_structure(doc_path: str) -> dict:
    """
    DOCX 파일에서 모든 테이블과 셀의 구조를 추출
    
    Args:
        doc_path: DOCX 파일 경로 (한글/공백 포함 가능)
    
    Returns:
        dict: 테이블 구조 정보
    """
    # 경로 정규화 (Windows 경로 처리)
    doc_path = os.path.normpath(doc_path)
    
    # 파일 존재 확인
    if not os.path.exists(doc_path):
        return {
            "error": f"File not found at '{doc_path}'"
        }
    
    # 파일 읽기 권한 확인
    if not os.access(doc_path, os.R_OK):
        return {
            "error": f"File is not readable: '{doc_path}'"
        }
    
    try:
        # 절대 경로로 변환하여 경로 문제 해결
        abs_path = os.path.abspath(doc_path)
        
        # 파일 크기 확인 (0바이트 파일 체크)
        file_size = os.path.getsize(abs_path)
        if file_size == 0:
            return {
                "error": f"File is empty (0 bytes): '{doc_path}'"
            }
        
        # python-docx로 파일 열기
        # Windows에서 한글 경로를 처리하기 위해 파일을 직접 열어서 전달
        # zipfile을 사용하여 DOCX 파일이 유효한지 먼저 확인
        import zipfile
        try:
            with zipfile.ZipFile(abs_path, 'r') as zip_ref:
                # DOCX는 ZIP 파일이므로 먼저 ZIP으로 열 수 있는지 확인
                pass
        except zipfile.BadZipFile:
            return {
                "error": f"File is not a valid DOCX file (not a valid ZIP archive): '{doc_path}'"
            }
        except Exception as zip_error:
            return {
                "error": f"Failed to open file as ZIP (DOCX files are ZIP archives): '{doc_path}'. Error: {str(zip_error)}"
            }
        
        # 이제 python-docx로 열기
        doc = Document(abs_path)
    except FileNotFoundError:
        return {
            "error": f"File not found: '{doc_path}' (absolute path: '{os.path.abspath(doc_path)}')"
        }
    except PermissionError:
        return {
            "error": f"Permission denied: '{doc_path}'"
        }
    except Exception as e:
        error_detail = str(e)
        error_type = type(e).__name__
        
        # "Package not found" 에러는 파일이 손상되었거나 DOCX 형식이 아닐 수 있음
        if "Package not found" in error_detail or "BadZipFile" in error_type:
            return {
                "error": f"Failed to open DOCX file. The file may be corrupted, not a valid DOCX file, or the path contains special characters that cannot be handled. Path: '{doc_path}' (absolute: '{os.path.abspath(doc_path)}'). Error: {error_detail}"
            }
        return {
            "error": f"Failed to open DOCX file '{doc_path}': {error_detail} (Type: {error_type})"
        }
    
    result = {
        "file_path": doc_path,
        "total_tables": 0,
        "tables": []
    }
    
    # 문서의 모든 요소를 순회하면서 테이블 찾기
    table_index = 0
    
    for element in doc.element.body:
        if isinstance(element, CT_Tbl):
            table = Table(element, doc)
            table_data = extract_table_data(table, table_index)
            result["tables"].append(table_data)
            table_index += 1
    
    result["total_tables"] = table_index
    
    return result


def extract_images_from_cell(cell) -> list:
    """
    셀에서 이미지 추출
    
    Args:
        cell: docx Table Cell 객체
        
    Returns:
        list: 이미지 정보 리스트
    """
    images = []
    found_ids = set()
    
    # paragraph의 runs에서 이미지 찾기
    for para in cell.paragraphs:
        for run in para.runs:
            # run.element에서 이미지 찾기
            drawings = run.element.findall('.//a:blip', namespaces={'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'})
            for drawing in drawings:
                embed_id = drawing.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                link_id = drawing.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}link')
                img_id = embed_id or link_id
                if img_id and img_id not in found_ids:
                    found_ids.add(img_id)
                    images.append({
                        "embed_id": embed_id,
                        "link_id": link_id,
                        "has_image": True
                    })
    
    return images


def extract_table_data(table: Table, table_index: int) -> dict:
    """
    테이블에서 모든 행과 셀 데이터 추출
    
    Args:
        table: docx Table 객체
        table_index: 테이블 인덱스
        
    Returns:
        dict: 테이블 데이터
    """
    table_data = {
        "table_index": table_index,
        "row_count": len(table.rows),
        "rows": []
    }
    
    for row_idx, row in enumerate(table.rows):
        row_data = {
            "row_index": row_idx,
            "cell_count": len(row.cells),
            "cells": []
        }
        
        for cell_idx, cell in enumerate(row.cells):
            # 셀의 모든 텍스트 추출 (여러 단락이 있을 수 있음)
            cell_text = "\n".join([para.text for para in cell.paragraphs])
            cell_text = cell_text.strip()
            
            # 셀 내 이미지 추출
            images = extract_images_from_cell(cell)
            
            cell_data = {
                "cell_index": cell_idx,
                "text": cell_text,
                "is_empty": len(cell_text) == 0 and len(images) == 0,
                "has_image": len(images) > 0,
                "image_count": len(images),
                "images": images,
                "position": {
                    "table_index": table_index,
                    "row_index": row_idx,
                    "cell_index": cell_idx
                }
            }
            
            row_data["cells"].append(cell_data)
        
        table_data["rows"].append(row_data)
    
    return table_data


def print_structure_summary(structure: dict, file=sys.stdout):
    """
    구조 정보를 읽기 쉬운 형식으로 출력
    """
    print("=" * 80, file=file)
    print(f"이력서 양식 구조 분석: {structure['file_path']}", file=file)
    print("=" * 80, file=file)
    print(f"\n총 테이블 개수: {structure['total_tables']}\n", file=file)
    
    for table in structure["tables"]:
        print(f"\n{'=' * 80}", file=file)
        print(f"테이블 {table['table_index']} (총 {table['row_count']}행)", file=file)
        print(f"{'=' * 80}", file=file)
        
        for row in table["rows"]:
            print(f"\n  행 {row['row_index']} (총 {row['cell_count']}개 셀):", file=file)
            
            for cell in row["cells"]:
                status_parts = []
                if cell["is_empty"]:
                    status_parts.append("빈칸")
                else:
                    if cell["text"]:
                        status_parts.append("텍스트 있음")
                    if cell["has_image"]:
                        status_parts.append(f"이미지 {cell['image_count']}개")
                status = " / ".join(status_parts) if status_parts else "빈칸"
                
                text_preview = cell["text"][:50] + "..." if len(cell["text"]) > 50 else cell["text"]
                
                print(f"    셀 [{row['row_index']}, {cell['cell_index']}]: {status}", file=file)
                if cell["text"]:
                    print(f"      텍스트: {text_preview}", file=file)
                if cell["has_image"]:
                    print(f"      이미지: {cell['image_count']}개 발견 (embed_id: {cell['images'][0]['embed_id'] if cell['images'] else 'N/A'})", file=file)
                print(f"      위치: 테이블{cell['position']['table_index']}, 행{cell['position']['row_index']}, 열{cell['position']['cell_index']}", file=file)


def print_mapping_template(structure: dict):
    """
    매핑 설정을 위한 템플릿 출력
    """
    print("\n\n" + "=" * 80)
    print("매핑 설정 템플릿 (resumeMapping.ts용)")
    print("=" * 80)
    print("\n다음 형식으로 각 칸에 어떤 데이터가 들어가야 하는지 알려주세요:\n")
    
    for table in structure["tables"]:
        print(f"\n// 테이블 {table['table_index']}")
        
        # 헤더 행 찾기 (첫 번째 행이 보통 헤더)
        if table["rows"]:
            header_row = table["rows"][0]
            print(f"// 헤더 행 {header_row['row_index']}:")
            for cell in header_row["cells"]:
                if cell["text"]:
                    print(f"//   열 {cell['cell_index']}: '{cell['text']}'")
        
        # 데이터 행들
        if len(table["rows"]) > 1:
            print(f"// 데이터 행: {table['rows'][1]['row_index']} ~ {table['rows'][-1]['row_index']}")
            print(f"// 각 행의 셀 내용:")
            for row in table["rows"][1:]:  # 헤더 제외
                print(f"//   행 {row['row_index']}:")
                for cell in row["cells"]:
                    print(f"//     열 {cell['cell_index']}: '{cell['text']}' (빈칸일 수 있음)")


def save_json_output(structure: dict, output_path: str = "resume_form_structure.json"):
    """
    구조 정보를 JSON 파일로 저장
    """
    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(structure, f, ensure_ascii=False, indent=2)
    print(f"\n\n구조 정보가 {output_path}에 저장되었습니다.")


def main():
    if len(sys.argv) < 2:
        error_msg = json.dumps({"error": "Usage: python3 scripts/extract_resume_form_structure.py <docx_file>"})
        print(error_msg, file=sys.stderr)
        sys.exit(1)
    
    # Windows에서 한글 경로 처리
    # sys.argv는 이미 올바른 인코딩으로 받아오지만, 안전을 위해 명시적으로 처리
    docx_path = sys.argv[1]
    
    # 경로 정규화 및 절대 경로 변환
    # Windows에서 한글 경로를 올바르게 처리하기 위해
    try:
        # 먼저 경로를 정규화
        docx_path = os.path.normpath(docx_path)
        # 절대 경로로 변환
        docx_path = os.path.abspath(docx_path)
    except Exception as e:
        error_msg = json.dumps({"error": f"Failed to process file path: {str(e)}"})
        print(error_msg, file=sys.stderr)
        sys.exit(1)
    
    if not os.path.exists(docx_path):
        error_msg = json.dumps({"error": f"File not found at '{docx_path}'"})
        print(error_msg, file=sys.stderr)
        sys.exit(1)
    
    try:
        # 구조 추출
        structure = extract_table_structure(docx_path)
        
        # 에러가 있으면 stderr로 출력하고 종료
        if "error" in structure:
            error_msg = json.dumps(structure)
            print(error_msg, file=sys.stderr)
            sys.exit(1)
        
        # JSON만 출력 (stdout으로 출력하여 Node.js에서 파싱 가능하도록)
        # stderr로는 요약 정보 출력 (선택적)
        json_output = json.dumps(structure, ensure_ascii=False, indent=2)
        print(json_output)
        
        # stderr로 요약 정보 출력 (선택적, 디버깅용)
        if os.getenv('VERBOSE') == '1':
            print_structure_summary(structure, file=sys.stderr)
        
    except Exception as e:
        error_msg = json.dumps({"error": f"Unexpected error: {str(e)}"})
        print(error_msg, file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()
